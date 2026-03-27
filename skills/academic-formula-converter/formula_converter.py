#!/usr/bin/env python3
"""
学术论文公式转换器 - 将Markdown中的LaTeX公式转换为docx和html格式
"""

import re
import os
import sys
import argparse

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import markdown
except ImportError:
    print("请安装依赖: pip install python-docx markdown")
    sys.exit(1)


class FormulaConverter:
    """学术论文公式转换器"""
    
    LATEX_TO_UNICODE = {
        r'\times': '×', r'\div': '÷', r'\pm': '±',
        r'\leq': '≤', r'\geq': '≥', r'\neq': '≠',
        r'\approx': '≈', r'\infty': '∞',
        r'\alpha': 'α', r'\beta': 'β', r'\gamma': 'γ',
        r'\delta': 'δ', r'\epsilon': 'ε', r'\theta': 'θ',
        r'\lambda': 'λ', r'\mu': 'μ', r'\pi': 'π',
        r'\sigma': 'σ', r'\phi': 'φ', r'\omega': 'ω',
        r'\sum': 'Σ', r'\prod': 'Π', r'\int': '∫',
        r'\partial': '∂', r'\nabla': '∇', r'\sqrt': '√',
        r'\rightarrow': '→', r'\leftarrow': '←',
    }
    
    def __init__(self, images_dir=None):
        self.images_dir = images_dir or './images'
        
    def convert_latex_to_unicode(self, formula):
        result = formula
        # 处理\frac
        while r'\frac' in result:
            result = re.sub(r'\\frac\{([^{}]*)\}\{([^{}]*)\}', r'((\1)/(\2))', result, count=1)
        # 处理其他命令
        for latex, uni in self.LATEX_TO_UNICODE.items():
            result = result.replace(latex, uni)
        # 清理
        result = re.sub(r'\^(\w|\d+)', r'^\1', result)
        result = re.sub(r'_\{?(\w|\d+)\}?', r'_\1', result)
        result = result.replace('{', '').replace('}', '').replace('\\', '')
        return result
    
    def process_markdown(self, content):
        # 块公式
        content = re.sub(r'\$\$([^$]+)\$\$', 
            lambda m: f"\n**{self.convert_latex_to_unicode(m.group(1))}**\n", content)
        # 行内公式
        content = re.sub(r'\$([^$]+)\$', 
            lambda m: self.convert_latex_to_unicode(m.group(1)), content)
        return content
    
    def convert_to_docx(self, md_path, docx_path, title=None):
        with open(md_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        md_content = self.process_markdown(md_content)
        doc = Document()
        
        for line in md_content.split('\n'):
            line = line.strip()
            if not line:
                continue
            
            # 图片
            img_match = re.search(r'!\[([^\]]*)\]\(([^\)]+)\)', line)
            if img_match and os.path.exists(img_match.group(2)):
                try:
                    doc.add_picture(img_match.group(2), width=Cm(14))
                    doc.add_paragraph(img_match.group(1)).alignment = WD_ALIGN_PARAGRAPH.CENTER
                except: pass
                continue
            
            # 标题
            if line.startswith('# '):
                doc.add_heading(level=1, text=line[2:])
            elif line.startswith('## '):
                doc.add_heading(level=2, text=line[3:])
            elif line.startswith('### '):
                doc.add_heading(level=3, text=line[3:])
            # 公式
            elif '**' in line:
                p = doc.add_paragraph()
                p.add_run(line.replace('**', '')).bold = True
            # 列表
            elif line.startswith('- '):
                doc.add_paragraph(line, style='List Bullet')
            else:
                doc.add_paragraph(line)
        
        doc.save(docx_path)
        print(f"已保存: {docx_path}")
        return True
    
    def convert_to_html(self, md_path, html_path, title=None):
        with open(md_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        md_content = self.process_markdown(md_content)
        html = markdown.markdown(md_content, extensions=['tables'])
        
        if not title:
            m = re.search(r'#\s+(.+)$', md_content, re.MULTILINE)
            title = m.group(1) if m else '学术论文'
        
        html_content = f'''<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>{title}</title>
<style>
body{{font-family:"SimSun",serif;line-height:1.8;padding:40px;max-width:900px;margin:0 auto}}
h1{{text-align:center;font-size:22pt}}h2{{font-size:16pt;border-bottom:1px solid #ccc}}
h3{{font-size:14pt}}p{{text-align:justify}}img{{max-width:100%;display:block;margin:20px auto}}
table{{border-collapse:collapse;width:100%}}th,td{{border:1px solid #333;padding:8px;text-align:center}}
</style></head><body>{html}</body></html>'''
        
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        print(f"已保存: {html_path}")
        return True
    
    def convert(self, input_path, output_path, output_format=None):
        if output_format is None:
            output_format = 'docx' if output_path.endswith('.docx') else 'html'
        
        title = None
        with open(input_path, 'r', encoding='utf-8') as f:
            for line in f:
                if line.startswith('# '):
                    title = line[2:].strip()
                    break
        
        if output_format == 'docx':
            return self.convert_to_docx(input_path, output_path, title)
        return self.convert_to_html(input_path, output_path, title)


def main():
    parser = argparse.ArgumentParser(description='学术论文公式转换器')
    parser.add_argument('input', help='输入的Markdown文件')
    parser.add_argument('output', help='输出的文件路径')
    parser.add_argument('--images', default='./images', help='图片目录')
    parser.add_argument('--format', choices=['docx', 'html'], help='输出格式')
    args = parser.parse_args()
    
    converter = FormulaConverter(images_dir=args.images)
    converter.convert(args.input, args.output, args.format)


if __name__ == '__main__':
    main()
